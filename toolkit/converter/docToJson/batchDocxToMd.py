#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Batch DOCX to Markdown Converter
Scans a folder for DOCX files and converts them to Markdown
"""

import os
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn


def docx_to_markdown(docx_path: str) -> str:
    """Convert a DOCX file to Markdown string."""
    doc = Document(docx_path)
    markdown_lines = []
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            markdown_lines.append("")
            continue
        
        # Determine heading level based on paragraph style
        style_name = para.style.name.lower() if para.style else ''
        
        if style_name.startswith('heading'):
            level = int(style_name[-1]) if style_name[-1].isdigit() else 1
            markdown_lines.append(f"{'#' * level} {text}")
        elif para.style.name == 'Title':
            markdown_lines.append(f"# {text}")
        else:
            markdown_lines.append(text)
    
    # Handle tables
    for table in doc.tables:
        markdown_lines.append("")
        # Table header
        header_cells = [cell.text.strip() for cell in table.rows[0].cells]
        markdown_lines.append("| " + " | ".join(header_cells) + " |")
        markdown_lines.append("| " + " | ".join(["---"] * len(header_cells)) + " |")
        
        # Table rows
        for row in table.rows[1:]:
            row_cells = [cell.text.strip() for cell in row.cells]
            markdown_lines.append("| " + " | ".join(row_cells) + " |")
        markdown_lines.append("")
    
    return "\n".join(markdown_lines)


def convert_folder(source_dir: str, target_dir: str) -> None:
    """
    Scan source folder for DOCX files and convert them to Markdown.
    
    Args:
        source_dir: Directory containing DOCX files
        target_dir: Directory to save Markdown files
    """
    source_path = Path(source_dir)
    target_path = Path(target_dir)
    
    # Create target directory if it doesn't exist
    target_path.mkdir(parents=True, exist_ok=True)
    
    # Find all DOCX files
    docx_files = list(source_path.glob("*.docx"))
    
    if not docx_files:
        print(f"No DOCX files found in {source_dir}")
        return
    
    print(f"Found {len(docx_files)} DOCX file(s)")
    
    # Convert each DOCX file
    success_count = 0
    error_count = 0
    
    for docx_file in docx_files:
        try:
            print(f"Converting: {docx_file.name}")
            
            # Convert to markdown
            markdown_content = docx_to_markdown(str(docx_file))
            
            # Generate output filename (same name but .md extension)
            md_filename = docx_file.stem + ".md"
            md_path = target_path / md_filename
            
            # Write markdown file
            with open(md_path, 'w', encoding='utf-8') as f:
                f.write(markdown_content)
            
            print(f"  ✓ Saved to: {md_path}")
            success_count += 1
            
        except Exception as e:
            print(f"  ✗ Error converting {docx_file.name}: {str(e)}")
            error_count += 1
    
    print(f"\nConversion complete: {success_count} succeeded, {error_count} failed")


if __name__ == "__main__":
    SOURCE_DIR = "/workspace/zbackup/sources"
    TARGET_DIR = "/workspace/dlake/edu/english/pschool"
    
    convert_folder(SOURCE_DIR, TARGET_DIR)
